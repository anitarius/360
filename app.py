import streamlit as st
import pandas as pd
import io
from PIL import Image
from docx import Document
import docx


def frontend():
    image_path = "logo.jpeg"  # Debe estar en la misma carpeta que app.py
    try:
        portada = Image.open(image_path)
        st.image(portada, use_column_width=True)
    except FileNotFoundError:
        st.warning(" No se encontró la imagen 'portada.jpg'. Asegurate de que esté en la misma carpeta que app.py.")

    st.set_page_config(page_title="360° AUTO", layout="centered")
    st.title("360° AUTO")
    st.markdown(
        """
        1) Subir el archivo con el que se quiere crear el reporte: excel con el formato prestablecido.
        2) Descargar informe procesado.
        """
    )
    nombre = st.text_input("Nombre del evaluado", value="Marina")
    apellido = st.text_input("Apellido del evaluado", value="Garrido")
    uploaded_file = st.file_uploader("Subir archivo de Excel", type=["xlsx"])
   


    
    if uploaded_file is not None:
        entrada = uploaded_file.name
        
        st.write(f"Archivo subido: {entrada}")

        procesar_archivo(uploaded_file,nombre,apellido)


def procesar_archivo(uploaded_file,nombre,apellido):
    # Cargar plantilla a modificar

    doc = docx.Document('plantilla.docx')

    # Carga y prerocesamiento del archivo
    df = pd.read_excel(uploaded_file)
    print(df.columns)

    #Comenzar a modificar el domuento
    for paragraph in doc.paragraphs:
        if 'XXX' in paragraph.text:
            paragraph.text = paragraph.text.replace('XXX', f'{nombre.upper()} {apellido.upper()}')
        if 'XX' in paragraph.text:
            paragraph.text = paragraph.text.replace('XX', f'{nombre}')

    #Comenzar a editar tablas
    diccionario_respuestas = {}
    preguntas = df.columns[3:]  # Asumiendo que las preguntas empiezan en la columna 3

    # Inicializa los diccionarios para cada pregunta
    for pregunta in preguntas:
        diccionario_respuestas[pregunta] = {}

    # Llena los diccionarios con las respuestas por categoría, agregando un asterisco al inicio de cada respuesta
    for pregunta in preguntas:
        categorias = df['categoria'].unique()  # Reemplaza 'categoria' con el nombre de la columna que contiene las categorías
        for categoria in categorias:
            respuestas = df[(df[pregunta].notna()) & (df['categoria'] == categoria)][pregunta].tolist()
            # Agrega un asterisco al principio de cada respuesta
            respuestas_con_asterisco = [f"* {respuesta}" for respuesta in respuestas]
            diccionario_respuestas[pregunta][categoria] = respuestas_con_asterisco



        import random

    tabla_index = 0
    preguntas = list(diccionario_respuestas.keys())  # Extraemos las preguntas del diccionario

    # Recorre las tablas del documento
    for tabla in doc.tables:
        if tabla_index == 0:
            # En la primera tabla, agrega las direcciones de correo electrónico
            if len(tabla.rows) > 1:
                direccion_correo_columna = df['Dirección de correo electrónico'].tolist()
                tabla.rows[1].cells[0].text = '\n'.join(direccion_correo_columna)
        else:
            # En las demás tablas, asigna una pregunta específica a cada tabla
            pregunta = preguntas[tabla_index - 1]  # Selecciona la pregunta basada en el índice de la tabla
            if pregunta:
                categorias = diccionario_respuestas[pregunta]
                fila_actual = 0
                if fila_actual + 2 < len(tabla.rows):
                    tabla.rows[fila_actual].cells[0].text = pregunta
                    fila_actual += 1
                    for categoria, respuestas in categorias.items():
                        if fila_actual + 1 < len(tabla.rows):
                            random.shuffle(respuestas)  # Mezcla las respuestas al azar
                            tabla.rows[fila_actual].cells[0].text = categoria
                            tabla.rows[fila_actual + 1].cells[0].text = '\n'.join(respuestas)
                            fila_actual += 2
                        else:
                            break
        tabla_index += 1

    #Termino de armar el archivo
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Aplica formato a la primera fila de cada tabla del documento
    for tabla in doc.tables:
        fila = tabla.rows[0]
        for paragraph in fila.cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)  # Blanco
            # Relleno azul para la celda
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '1F497D')  # Azul
            fila.cells[0]._element.get_or_add_tcPr().append(shading_elm)

    # Recorre las tablas (empezando desde la segunda tabla)
    for tabla_index, tabla in enumerate(doc.tables[1:], start=1):
        for fila_index, fila in enumerate(tabla.rows):
            texto_celda = fila.cells[0].text.strip()

            # Verifica si el texto de la celda es una pregunta
            if texto_celda in diccionario_respuestas.keys():
                # Aplica el formato a la celda de pregunta
                for paragraph in fila.cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(14)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)  # Blanco
                    # Relleno azul para la celda
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), '1F497D')  # Azul
                    fila.cells[0]._element.get_or_add_tcPr().append(shading_elm)

            # Aplica formato a la segunda fila de la primera tabla solamente
            elif tabla_index == 1 and fila_index == 1:
                for paragraph in fila.cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(14)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0, 112, 192)  # Azul

            # Verifica si el texto de la celda es una categoría
            elif any(texto_celda in categorias for categorias in diccionario_respuestas.values()):
                # Aplica el formato a la celda de categoría
                for paragraph in fila.cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(14)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0, 112, 192)  # Azul
                        run.text = run.text.upper()  # Mayúsculas

            # Si es una respuesta
            else:
                for paragraph in fila.cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(0, 0, 0)  # Negro

    #salida
    salida = f'Reporte 360 {nombre} {apellido}.docx'

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="⬇️ Descargar informe Word",
        data=buffer,
        file_name=f"{salida}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

frontend()
